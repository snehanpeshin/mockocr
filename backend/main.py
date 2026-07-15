from __future__ import annotations

from html import escape
import shutil
import tempfile
import os
from pathlib import Path

import cv2
from docx import Document
from fastapi import FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import numpy as np
from pydantic import BaseModel

from image_preprocess import preprocess_image_variants
from firebase_auth import require_matching_email
from beta_service import (
    beta_summary,
    feedback_summary,
    request_tablet_preorder,
    request_beta_access,
    save_customer_discovery,
    tablet_preorder_summary,
    verify_beta_token,
)
from cost_control import (
    DuplicateInProgress,
    LimitExceeded,
    MonetizationRequired,
    ServiceDisabled,
    abandon_cache,
    admin_usage_summary,
    cache_key_for,
    client_status,
    complete_cache,
    effective_access,
    enforce_kill_switch,
    identify_scan_request,
    max_pages_per_upload,
    max_upload_bytes,
    reserve_limits,
    reserve_or_get_cache,
    usage_event,
)
from note_service import delete_note, save_note, search_notes
from ocr_service import clean_ocr_text, extract_text
from payment_service import (
    construct_webhook_event,
    create_checkout_session,
    create_payment_link,
    handle_stripe_event,
    revenue_summary,
    validate_admin_token,
)
from scan_event_service import record_scan_event, scan_summary


OUTPUTS_DIR = Path(__file__).resolve().parent.parent / "outputs"
DEFAULT_FRONTEND_ORIGINS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://main.d3vhgcrptn13ws.amplifyapp.com",
    "https://mockocr.com",
    "https://www.mockocr.com",
    "https://cleanote.in",
    "https://www.cleanote.in",
]


def _frontend_origins() -> list[str]:
    configured_origins = os.getenv("FRONTEND_ORIGINS", "")
    origins = DEFAULT_FRONTEND_ORIGINS + [
        origin.strip() for origin in configured_origins.split(",") if origin.strip()
    ]
    return list(dict.fromkeys(origins))

app = FastAPI(title="Cleanote API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=_frontend_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ExportRequest(BaseModel):
    text: str


class BetaRequest(BaseModel):
    name: str
    email: str
    role: str


class TabletPreorderRequest(BaseModel):
    name: str
    email: str
    role: str
    quantity: int = 1
    use_case: str = ""


class DiscoveryRequest(BaseModel):
    email: str
    name: str = ""
    role: str = ""
    source: str = "post_scan"
    note_filename: str = ""
    subject: str = ""
    word_count: int = 0
    rating: int = 0
    feedback: str = ""
    worked: str = ""
    missing: str = ""
    pay_value: str = ""


class NoteRequest(BaseModel):
    id: str
    email: str
    createdAt: str
    filename: str
    provider: str
    subject: str
    text: str
    contextText: str = ""


class CheckoutRequest(BaseModel):
    product_key: str
    customer_email: str | None = None
    success_url: str
    cancel_url: str


class PaymentLinkRequest(BaseModel):
    product_key: str


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/")
def root() -> dict[str, str]:
    return {"status": "ok", "service": "cleanote-backend"}


@app.post("/api/beta/request")
def request_beta(payload: BetaRequest) -> dict[str, object]:
    try:
        return request_beta_access(payload.name, payload.email, payload.role)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/api/beta/verify")
def verify_beta(token: str) -> dict[str, object]:
    try:
        return verify_beta_token(token)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except TimeoutError as exc:
        raise HTTPException(status_code=410, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/beta/discovery")
def save_beta_discovery(payload: DiscoveryRequest) -> dict[str, str]:
    try:
        return save_customer_discovery(payload.model_dump())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/tablet/preorder")
def request_tablet_bundle_preorder(payload: TabletPreorderRequest) -> dict[str, str]:
    try:
        return request_tablet_preorder(payload.model_dump())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/notes")
def save_user_note(
    payload: NoteRequest,
    authorization: str | None = Header(default=None),
) -> dict[str, str]:
    try:
        require_matching_email(authorization, payload.email)
        return save_note(payload.model_dump())
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/api/notes/search")
def search_user_notes(
    email: str,
    q: str = "",
    limit: int = 30,
    authorization: str | None = Header(default=None),
) -> dict[str, object]:
    try:
        require_matching_email(authorization, email)
        return search_notes(email, q, limit)
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.delete("/api/notes/{note_id}")
def delete_user_note(
    note_id: str,
    email: str,
    authorization: str | None = Header(default=None),
) -> dict[str, str]:
    try:
        require_matching_email(authorization, email)
        return delete_note(email, note_id)
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/stripe/checkout-session")
def create_stripe_checkout_session(payload: CheckoutRequest) -> dict[str, str]:
    try:
        return create_checkout_session(
            payload.product_key,
            payload.customer_email,
            payload.success_url,
            payload.cancel_url,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/stripe/payment-link")
def create_stripe_payment_link(
    payload: PaymentLinkRequest,
    x_admin_token: str | None = Header(default=None),
) -> dict[str, str]:
    try:
        validate_admin_token(x_admin_token)
        return create_payment_link(payload.product_key)
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/stripe/webhook")
async def stripe_webhook(request: Request, stripe_signature: str | None = Header(default=None)):
    payload = await request.body()
    try:
        event = construct_webhook_event(payload, stripe_signature)
        return handle_stripe_event(event)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/admin/revenue")
def admin_revenue(x_admin_token: str | None = Header(default=None)) -> dict[str, object]:
    try:
        validate_admin_token(x_admin_token)
        return {
            **revenue_summary(),
            "beta_summary": _safe_beta_summary(),
            "scan_summary": scan_summary(),
            "feedback_summary": feedback_summary(),
            "tablet_preorder_summary": _safe_tablet_preorder_summary(),
        }
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/api/admin/usage")
def admin_usage(
    days: int = 1,
    x_admin_token: str | None = Header(default=None),
) -> dict[str, object]:
    try:
        validate_admin_token(x_admin_token)
        return admin_usage_summary(days)
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/api/config/status")
def config_status(
    authorization: str | None = Header(default=None),
    x_cleanote_installation_id: str | None = Header(default=None),
) -> dict[str, object]:
    try:
        identity = identify_scan_request(authorization, x_cleanote_installation_id)
    except Exception:
        identity = None
    return client_status(identity)


def _safe_beta_summary() -> dict[str, object]:
    try:
        return beta_summary()
    except Exception as exc:
        return {
            "available": False,
            "error": str(exc),
            "signup_count": 0,
            "beta_access_count": 0,
            "manual_required_count": 0,
            "emailed_count": 0,
            "recent_signups": [],
        }


def _safe_tablet_preorder_summary() -> dict[str, object]:
    try:
        return tablet_preorder_summary()
    except Exception as exc:
        return {
            "available": False,
            "error": str(exc),
            "preorder_count": 0,
            "total_quantity": 0,
            "recent_preorders": [],
        }


@app.post("/api/ocr")
async def run_ocr(
    file: UploadFile = File(...),
    provider: str | None = Form(default=None),
    subject: str = Form(default="general"),
    context_text: str = Form(default=""),
    cleanup_mode: str = Form(default="rules"),
    authorization: str | None = Header(default=None),
    x_cleanote_installation_id: str | None = Header(default=None),
    idempotency_key: str | None = Header(default=None),
) -> dict[str, object]:
    filename = file.filename or "upload"
    suffix = Path(filename).suffix.lower()
    if suffix not in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".pdf", ".docx"}:
        record_scan_event(
            {
                "filename": filename,
                "file_type": suffix.lstrip("."),
                "provider": provider or "",
                "subject": subject,
                "status": "rejected",
                "error_message": "Unsupported file type",
            }
        )
        raise HTTPException(
            status_code=400,
            detail="Upload an image, PDF, or DOCX file.",
        )

    upload_bytes = await file.read()
    file_size_bytes = len(upload_bytes)
    if file_size_bytes > max_upload_bytes():
        record_scan_event(
            {
                "filename": filename,
                "file_type": suffix.lstrip("."),
                "file_size_bytes": file_size_bytes,
                "provider": provider or "",
                "subject": subject,
                "status": "rejected",
                "error_message": "Upload exceeds server size limit.",
            }
        )
        raise HTTPException(status_code=413, detail="This upload is too large. Try fewer pages or a smaller image.")

    cache_reservation = None
    try:
        identity = identify_scan_request(authorization, x_cleanote_installation_id)
        access = effective_access(identity, cleanup_mode)
        enforce_kill_switch(identity, access)

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            uploaded_path = temp_path / f"upload{suffix}"
            original_path = temp_path / "original.png" if suffix == ".pdf" else uploaded_path

            uploaded_path.write_bytes(upload_bytes)

            if suffix == ".docx":
                extracted_text = clean_ocr_text(_extract_docx_text(uploaded_path))
                record_scan_event(usage_event(
                    identity=identity,
                    access=access,
                    filename=filename,
                    file_type=suffix.lstrip("."),
                    upload_bytes=file_size_bytes,
                    page_count=1,
                    provider="docx",
                    status="success",
                    cache_hit=False,
                    idempotency_key=idempotency_key,
                    subject=subject,
                    text_length=len(extracted_text),
                ))

                return {
                    "text": extracted_text,
                    "provider": "docx",
                    "filename": filename,
                    "subject": subject,
                    "context_text": context_text,
                    "usage_tier": access.tier,
                    "effective_cleanup_mode": access.cleanup_mode,
                    "cache_hit": False,
                }

            page_paths = (
                _render_pdf_pages(uploaded_path, temp_path / "pdf_pages")
                if suffix == ".pdf"
                else [original_path]
            )
            page_count = len(page_paths)
            if page_count > max_pages_per_upload():
                record_scan_event(usage_event(
                    identity=identity,
                    access=access,
                    filename=filename,
                    file_type=suffix.lstrip("."),
                    upload_bytes=file_size_bytes,
                    page_count=page_count,
                    provider=provider or "",
                    status="rejected",
                    cache_hit=False,
                    idempotency_key=idempotency_key,
                    subject=subject,
                    error_message="Upload exceeds page limit.",
                ))
                raise HTTPException(
                    status_code=413,
                    detail=f"This upload has {page_count} pages. Upload {max_pages_per_upload()} pages or fewer.",
                )

            cache_reservation = cache_key_for(
                identity=identity,
                upload_bytes=upload_bytes,
                provider=provider or os.getenv("OCR_PROVIDER", "mock"),
                cleanup_mode=access.cleanup_mode,
                subject=subject,
                options={
                    "context_text_hash": context_text,
                    "fast_multi_page": page_count > 1,
                    "file_type": suffix.lstrip("."),
                },
            )
            cache_reservation = reserve_or_get_cache(cache_reservation, idempotency_key)
            if cache_reservation.cache_hit and cache_reservation.cached_response:
                cached_response = dict(cache_reservation.cached_response)
                record_scan_event(usage_event(
                    identity=identity,
                    access=access,
                    filename=filename,
                    file_type=suffix.lstrip("."),
                    upload_bytes=file_size_bytes,
                    page_count=page_count,
                    provider=str(cached_response.get("provider") or provider or ""),
                    status="success",
                    cache_hit=True,
                    idempotency_key=idempotency_key,
                    subject=subject,
                    text_length=len(str(cached_response.get("text") or "")),
                ))
                cached_response["cache_hit"] = True
                return cached_response

            reservation = reserve_limits(identity, access, page_count)
            page_results: list[dict[str, str]] = []
            fast_multi_page = page_count > 1
            for page_index, page_path in enumerate(page_paths):
                processed_paths = preprocess_image_variants(
                    page_path,
                    temp_path / f"ocr_candidates_page_{page_index + 1}",
                )
                if not processed_paths:
                    raise ValueError("Could not prepare this image for OCR.")
                page_results.append(
                    extract_text(
                        processed_paths,
                        provider,
                        subject,
                        context_text,
                        fast_mode=fast_multi_page,
                        cleanup_mode=access.cleanup_mode,
                    )
                )

            result_text = _combine_page_results(page_results)
            result_provider = _combine_providers(page_results)
            response_payload: dict[str, object] = {
                "text": result_text,
                "provider": result_provider,
                "filename": filename,
                "subject": subject,
                "context_text": context_text,
                "usage_tier": access.tier,
                "effective_cleanup_mode": access.cleanup_mode,
                "cache_hit": False,
            }
            complete_cache(cache_reservation, response_payload, idempotency_key)
            record_scan_event(usage_event(
                identity=identity,
                access=access,
                filename=filename,
                file_type=suffix.lstrip("."),
                upload_bytes=file_size_bytes,
                page_count=page_count,
                provider=result_provider,
                status="success",
                cache_hit=False,
                idempotency_key=idempotency_key,
                subject=subject,
                text_length=len(result_text),
            ))
            _ = reservation
            return response_payload
    except DuplicateInProgress as exc:
        raise HTTPException(status_code=409, detail=exc.detail) from exc
    except LimitExceeded as exc:
        abandon_cache(cache_reservation)
        record_scan_event(
            {
                "filename": filename,
                "file_type": suffix.lstrip("."),
                "file_size_bytes": file_size_bytes,
                "provider": provider or "",
                "subject": subject,
                "status": "rejected",
                "error_message": "Scan limit exceeded.",
            }
        )
        raise HTTPException(status_code=429, detail=exc.detail) from exc
    except MonetizationRequired as exc:
        raise HTTPException(status_code=402, detail=exc.detail) from exc
    except ServiceDisabled as exc:
        raise HTTPException(status_code=503, detail=exc.detail) from exc
    except HTTPException:
        raise
    except ValueError as exc:
        abandon_cache(cache_reservation)
        record_scan_event(
            {
                "filename": filename,
                "file_type": suffix.lstrip("."),
                "file_size_bytes": file_size_bytes,
                "provider": provider or "",
                "subject": subject,
                "status": "failed",
                "error_message": str(exc),
            }
        )
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        abandon_cache(cache_reservation)
        record_scan_event(
            {
                "filename": filename,
                "file_type": suffix.lstrip("."),
                "file_size_bytes": file_size_bytes,
                "provider": provider or "",
                "subject": subject,
                "status": "failed",
                "error_message": str(exc),
            }
        )
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/preview/docx")
async def preview_docx(file: UploadFile = File(...)) -> dict[str, str]:
    filename = file.filename or "upload.docx"
    if not filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Upload a DOCX file.")

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            uploaded_path = Path(temp_dir) / "preview.docx"
            with uploaded_path.open("wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            return {
                "filename": filename,
                "html": _docx_preview_html(uploaded_path),
                "text": clean_ocr_text(_extract_docx_text(uploaded_path)),
            }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


def _render_pdf_pages(pdf_path: Path, output_dir: Path) -> list[Path]:
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise RuntimeError("Install pypdfium2 to upload PDF files.") from exc

    document = pdfium.PdfDocument(pdf_path)
    if len(document) == 0:
        raise ValueError("The uploaded PDF has no pages.")

    output_dir.mkdir(parents=True, exist_ok=True)
    max_pages = max(1, int(os.getenv("OCR_MAX_PDF_PAGES", "5")))
    rendered_paths: list[Path] = []

    for page_index in range(min(len(document), max_pages)):
        page = document[page_index]
        output_path = output_dir / f"page_{page_index + 1}.png"
        bitmap = page.render(scale=3.0)
        image = np.asarray(bitmap.to_numpy())
        if image.shape[-1] == 4:
            image = cv2.cvtColor(image, cv2.COLOR_RGBA2BGR)
        else:
            image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
        cv2.imwrite(str(output_path), image)
        rendered_paths.append(output_path)

    return rendered_paths


def _combine_page_results(page_results: list[dict[str, str]]) -> str:
    if len(page_results) == 1:
        return page_results[0]["text"]

    return "\n\n---\n\n".join(
        f"Page {index + 1}\n\n{result['text']}" for index, result in enumerate(page_results)
    )


def _combine_providers(page_results: list[dict[str, str]]) -> str:
    providers = [result["provider"] for result in page_results if result.get("provider")]
    if not providers:
        return "unknown"
    unique_providers = list(dict.fromkeys(providers))
    if len(unique_providers) == 1:
        return unique_providers[0]
    return f"{unique_providers[0]}+{len(unique_providers)}pages"


def _extract_docx_text(docx_path: Path) -> str:
    document = Document(docx_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("The uploaded DOCX did not contain readable text.")
    return text


def _docx_preview_html(docx_path: Path) -> str:
    document = Document(docx_path)
    parts: list[str] = []

    for block in document.element.body:
        tag_name = block.tag.rsplit("}", 1)[-1]
        if tag_name == "p":
            paragraph = next(
                (
                    paragraph
                    for paragraph in document.paragraphs
                    if paragraph._element is block
                ),
                None,
            )
            if paragraph is not None:
                parts.append(_docx_paragraph_html(paragraph))
        elif tag_name == "tbl":
            table = next((table for table in document.tables if table._element is block), None)
            if table is not None:
                parts.append(_docx_table_html(table))

    html = "\n".join(part for part in parts if part).strip()
    if not html:
        raise ValueError("The uploaded DOCX did not contain readable preview content.")
    return html


def _docx_paragraph_html(paragraph) -> str:
    text_parts: list[str] = []
    for run in paragraph.runs:
        run_text = escape(run.text)
        if not run_text:
            continue
        if run.bold:
            run_text = f"<strong>{run_text}</strong>"
        if run.italic:
            run_text = f"<em>{run_text}</em>"
        if run.underline:
            run_text = f"<u>{run_text}</u>"
        text_parts.append(run_text)

    text = "".join(text_parts).strip()
    if not text:
        return ""

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "heading 1" in style_name or style_name == "title":
        tag = "h1"
    elif "heading 2" in style_name:
        tag = "h2"
    elif "heading" in style_name:
        tag = "h3"
    else:
        tag = "p"

    alignment = paragraph.alignment
    align_class = ""
    if alignment is not None:
        align_value = str(alignment).split(" ")[0].lower()
        if align_value in {"center", "right", "justify"}:
            align_class = f' class="docx-align-{align_value}"'

    return f"<{tag}{align_class}>{text}</{tag}>"


def _docx_table_html(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = "<br />".join(
                escape(paragraph.text.strip())
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            cells.append(f"<td>{cell_text}</td>")
        if cells:
            rows.append(f"<tr>{''.join(cells)}</tr>")

    if not rows:
        return ""
    return f"<table><tbody>{''.join(rows)}</tbody></table>"


@app.post("/api/export/txt")
def export_txt(payload: ExportRequest) -> FileResponse:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUTS_DIR / "cleanote-output.txt"
    output_path.write_text(payload.text, encoding="utf-8")
    return FileResponse(
        output_path,
        media_type="text/plain",
        filename="cleanote-output.txt",
    )


@app.post("/api/export/docx")
def export_docx(payload: ExportRequest) -> FileResponse:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUTS_DIR / "cleanote-output.docx"

    document = Document()
    for paragraph in payload.text.splitlines() or [""]:
        document.add_paragraph(paragraph)
    document.save(output_path)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="cleanote-output.docx",
    )


@app.post("/api/export/pdf")
def export_pdf(payload: ExportRequest) -> FileResponse:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUTS_DIR / "cleanote-output.pdf"

    output_path.write_bytes(_build_simple_pdf(payload.text))

    return FileResponse(
        output_path,
        media_type="application/pdf",
        filename="cleanote-output.pdf",
    )


def _build_simple_pdf(text: str) -> bytes:
    lines = _wrap_pdf_lines(text)
    pages = [lines[index : index + 48] for index in range(0, len(lines), 48)] or [[""]]
    font_object_id = 3 + (len(pages) * 2)
    objects = [b"<< /Type /Catalog /Pages 2 0 R >>"]

    page_object_ids = [3 + index for index in range(len(pages))]
    page_refs = " ".join(f"{object_id} 0 R" for object_id in page_object_ids)
    objects.append(f"<< /Type /Pages /Kids [{page_refs}] /Count {len(pages)} >>".encode("ascii"))

    content_objects: list[bytes] = []
    for page_index, page_lines in enumerate(pages):
        content_object_id = 3 + len(pages) + page_index
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_object_id} 0 R >> >> "
                f"/Contents {content_object_id} 0 R >>"
            ).encode("ascii")
        )
        stream = _pdf_text_stream(page_lines)
        content_objects.append(
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream"
        )

    objects.extend(content_objects)
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    return _pdf_from_objects(objects)


def _pdf_text_stream(lines: list[str]) -> bytes:
    content_lines = ["BT", "/F1 11 Tf", "14 TL", "48 744 Td"]
    for line in lines:
        content_lines.append(f"({_escape_pdf_string(line)}) Tj")
        content_lines.append("T*")
    content_lines.append("ET")
    return "\n".join(content_lines).encode("latin-1", errors="replace")


def _wrap_pdf_lines(text: str, width: int = 88) -> list[str]:
    source_lines = text.splitlines() or [""]
    wrapped: list[str] = []
    for source_line in source_lines:
        line = source_line.strip()
        if not line:
            wrapped.append("")
            continue
        while len(line) > width:
            split_at = line.rfind(" ", 0, width)
            if split_at < 24:
                split_at = width
            wrapped.append(line[:split_at].strip())
            line = line[split_at:].strip()
        wrapped.append(line)
    return wrapped


def _escape_pdf_string(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _pdf_from_objects(objects: list[bytes]) -> bytes:
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)
